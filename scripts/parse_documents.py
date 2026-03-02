"""
多格式文档解析器
支持 PDF、Word、Excel、Markdown、TXT 等格式
统一转换为项目的标准 JSON 格式
"""
import os
import json
from typing import Dict, List, Optional
from pathlib import Path
from datetime import datetime
import re

# PDF 解析
try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    print("警告: pypdf 未安装，无法解析 PDF。请运行: pip install pypdf")

# Word 解析
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: python-docx 未安装，无法解析 Word。请运行: pip install python-docx")

# Excel 解析
try:
    import pandas as pd
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False
    print("警告: pandas 未安装，无法解析 Excel。请运行: pip install pandas openpyxl")

# HTML 解析
try:
    from bs4 import BeautifulSoup
    HAS_HTML = True
except ImportError:
    HAS_HTML = False


class DocumentParser:
    """统一的文档解析器"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.parse_pdf,
            '.docx': self.parse_docx,
            '.doc': self.parse_docx,
            '.xlsx': self.parse_excel,
            '.xls': self.parse_excel,
            '.csv': self.parse_csv,
            '.txt': self.parse_txt,
            '.md': self.parse_markdown,
            '.html': self.parse_html,
            '.htm': self.parse_html
        }
    
    def parse(self, file_path: str, category: Optional[str] = None) -> Dict:
        """
        解析文档并转换为标准格式
        
        Args:
            file_path: 文档路径
            category: 文档分类（可选，如 HR、IT、Finance）
        
        Returns:
            标准格式的文档字典
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        if suffix not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {suffix}")
        
        # 调用对应的解析方法
        parse_func = self.supported_formats[suffix]
        content = parse_func(file_path)
        
        # 构建标准文档对象
        doc = {
            "doc_id": self._generate_doc_id(file_path),
            "title": file_path.stem,  # 文件名（不含扩展名）
            "category": category or "Imported",
            "source_type": f"imported_{suffix[1:]}",
            "source_file": str(file_path),
            "content": content,
            "word_count": len(content),
            "generated_at": datetime.now().isoformat(),
            "metadata": {
                "original_format": suffix,
                "file_size": file_path.stat().st_size
            }
        }
        
        return doc
    
    def parse_pdf(self, file_path: Path) -> str:
        """解析 PDF 文件"""
        if not HAS_PDF:
            raise ImportError("需要安装 pypdf: pip install pypdf")
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_parts.append(f"[第 {page_num} 页]\n{text}")
        
        return "\n\n".join(text_parts)
    
    def parse_docx(self, file_path: Path) -> str:
        """解析 Word 文档"""
        if not HAS_DOCX:
            raise ImportError("需要安装 python-docx: pip install python-docx")
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 解析表格
        for table in doc.tables:
            table_text = self._parse_word_table(table)
            if table_text:
                text_parts.append(f"\n【表格】\n{table_text}")
        
        return "\n\n".join(text_parts)
    
    def _parse_word_table(self, table) -> str:
        """解析 Word 表格"""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
    
    def parse_excel(self, file_path: Path) -> str:
        """解析 Excel 文件"""
        if not HAS_EXCEL:
            raise ImportError("需要安装 pandas 和 openpyxl: pip install pandas openpyxl")
        
        # 读取所有 sheet
        excel_file = pd.ExcelFile(file_path)
        text_parts = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # 转换为文本格式
            sheet_text = f"【Sheet: {sheet_name}】\n"
            sheet_text += df.to_string(index=False)
            text_parts.append(sheet_text)
        
        return "\n\n".join(text_parts)
    
    def parse_csv(self, file_path: Path) -> str:
        """解析 CSV 文件"""
        if not HAS_EXCEL:
            raise ImportError("需要安装 pandas: pip install pandas")
        
        df = pd.read_csv(file_path)
        return df.to_string(index=False)
    
    def parse_txt(self, file_path: Path) -> str:
        """解析纯文本文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def parse_markdown(self, file_path: Path) -> str:
        """解析 Markdown 文件（保留结构）"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # 可选：移除代码块中的语言标识
        content = re.sub(r'```[\w]*\n', '```\n', content)
        
        return content
    
    def parse_html(self, file_path: Path) -> str:
        """解析 HTML 文件"""
        if not HAS_HTML:
            # 简单的 HTML 标签移除
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            # 移除 script 和 style
            html_content = re.sub(r'<script.*?</script>', '', html_content, flags=re.DOTALL)
            html_content = re.sub(r'<style.*?</style>', '', html_content, flags=re.DOTALL)
            # 移除所有 HTML 标签
            text = re.sub(r'<[^>]+>', '', html_content)
            return text.strip()
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                soup = BeautifulSoup(f, 'html.parser')
            
            # 移除 script 和 style
            for script in soup(["script", "style"]):
                script.decompose()
            
            return soup.get_text(separator="\n").strip()
    
    def _generate_doc_id(self, file_path: Path) -> str:
        """生成文档 ID"""
        # 使用文件名的 hash 生成唯一 ID
        import hashlib
        file_hash = hashlib.md5(str(file_path).encode()).hexdigest()[:8]
        return f"doc_imported_{file_hash}"
    
    def batch_parse(
        self, 
        directory: str, 
        output_dir: str = "./data/docs",
        category_mapping: Optional[Dict[str, str]] = None
    ) -> List[Dict]:
        """
        批量解析目录下的所有文档
        
        Args:
            directory: 文档目录
            output_dir: 输出目录
            category_mapping: 文件名模式 -> 分类的映射
                例如: {"hr_": "HR", "it_": "IT"}
        
        Returns:
            解析后的文档列表
        """
        directory = Path(directory)
        os.makedirs(output_dir, exist_ok=True)
        
        all_docs = []
        
        # 遍历目录
        for file_path in directory.rglob("*"):
            if not file_path.is_file():
                continue
            
            suffix = file_path.suffix.lower()
            if suffix not in self.supported_formats:
                continue
            
            try:
                # 确定分类
                category = None
                if category_mapping:
                    for pattern, cat in category_mapping.items():
                        if pattern.lower() in file_path.name.lower():
                            category = cat
                            break
                
                # 解析文档
                doc = self.parse(file_path, category=category)
                all_docs.append(doc)
                
                # 保存单个文档
                output_path = Path(output_dir) / f"{doc['doc_id']}.json"
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(doc, f, ensure_ascii=False, indent=2)
                
                print(f"✅ 已解析: {file_path.name} -> {doc['doc_id']}")
                
            except Exception as e:
                print(f"❌ 解析失败 [{file_path.name}]: {e}")
        
        # 保存汇总
        summary_path = Path(output_dir) / "_imported_docs.json"
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(all_docs, f, ensure_ascii=False, indent=2)
        
        print(f"\n📊 共解析 {len(all_docs)} 份文档")
        return all_docs


def main():
    """命令行工具"""
    import argparse
    
    parser = argparse.ArgumentParser(description="解析多格式文档并转换为标准格式")
    parser.add_argument("input", help="输入文件或目录路径")
    parser.add_argument("-o", "--output", default="./data/docs", help="输出目录")
    parser.add_argument("-c", "--category", help="文档分类（如 HR, IT, Finance）")
    
    args = parser.parse_args()
    
    parser_obj = DocumentParser()
    input_path = Path(args.input)
    
    if input_path.is_file():
        # 解析单个文件
        doc = parser_obj.parse(input_path, category=args.category)
        
        # 保存
        output_dir = Path(args.output)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"{doc['doc_id']}.json"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(doc, f, ensure_ascii=False, indent=2)
        
        print(f"✅ 文档已保存: {output_path}")
        
    elif input_path.is_dir():
        # 批量解析目录
        parser_obj.batch_parse(input_path, output_dir=args.output)
    else:
        print(f"❌ 路径不存在: {input_path}")


if __name__ == "__main__":
    main()
